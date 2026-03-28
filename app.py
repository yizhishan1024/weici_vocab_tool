from __future__ import annotations

import io
import json
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, asdict
from datetime import datetime
import os
from pathlib import Path
import time
from typing import List
from urllib.parse import quote, unquote, urlparse

import requests
from flask import Flask, Response, jsonify, render_template, request, send_file

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
CACHE_PATH = DATA_DIR / "words_cache.json"
NOTEBOOK_DIR = DATA_DIR / "notebook"
NOTEBOOK_PATH = NOTEBOOK_DIR / "selected_words.json"

RAW_BASE = "https://raw.githubusercontent.com/1299172402/weici/master/docs/2"
PAGES_MEDIA_BASE = "https://1299172402.github.io/weici/media"
SOURCE_FILES = [f"weici_{ch}.md" for ch in "ABCDEFGHIJKLMNOPQRSTUVWXYZ"] + ["weici_phrase.md"]

app = Flask(__name__)


@dataclass
class WordEntry:
    word: str
    pos: str
    category: str
    ipa_uk: str
    ipa_us: str
    audio_uk: str
    audio_us: str
    first_zh: str
    first_en: str
    source_path: str


def clean_pos(pos: str) -> str:
    value = re.sub(r"\s*重难点词汇\s*", " ", pos or "")
    return re.sub(r"\s+", " ", value).strip()


def trim_meaning(text: str) -> str:
    return (text or "").strip().rstrip("：:;； ")


def aggregate_words(words: List[WordEntry]) -> list[dict]:
    grouped: dict[str, list[WordEntry]] = {}
    for item in words:
        grouped.setdefault(item.word, []).append(item)

    aggregated = []
    for word in sorted(grouped.keys(), key=str.lower):
        group = grouped[word]
        first_item = group[0]
        meanings = []
        seen = set()
        pos_parts = []

        for item in group:
            pos = clean_pos(item.pos)
            zh = trim_meaning(item.first_zh)
            en = trim_meaning(item.first_en)
            key = (pos, zh, en)
            if key in seen:
                continue
            seen.add(key)

            meanings.append(
                {
                    "pos": pos,
                    "zh": zh,
                    "en": en,
                }
            )
            if pos and pos not in pos_parts:
                pos_parts.append(pos)
            if len(meanings) >= 2:
                break

        ipa = next((x.ipa_us for x in group if x.ipa_us), "") or next((x.ipa_uk for x in group if x.ipa_uk), "")
        audio = next((x.audio_us for x in group if x.audio_us), "") or next((x.audio_uk for x in group if x.audio_uk), "")
        aggregated.append(
            {
                "word": word,
                "pos": " / ".join(pos_parts),
                "category": first_item.category,
                "ipa": ipa,
                "ipa_uk": next((x.ipa_uk for x in group if x.ipa_uk), ""),
                "ipa_us": next((x.ipa_us for x in group if x.ipa_us), ""),
                "audio": audio,
                "audio_uk": next((x.audio_uk for x in group if x.audio_uk), ""),
                "audio_us": next((x.audio_us for x in group if x.audio_us), ""),
                "meanings": meanings,
                "first_zh": meanings[0]["zh"] if meanings else "",
                "first_en": meanings[0]["en"] if meanings else "",
                "source_path": first_item.source_path,
            }
        )

    return aggregated


def build_export_meaning(item: dict, mode: str) -> str:
    if mode == "word":
        return ""

    meanings = item.get("meanings") or []
    if meanings:
        parts = []
        for meaning in meanings[:2]:
            pos = (meaning.get("pos") or "").strip()
            zh = trim_meaning(meaning.get("zh") or "")
            en = trim_meaning(meaning.get("en") or "")

            if mode == "en":
                body = en
            elif mode == "bilingual":
                body = f"{zh} / {en}" if zh and en else (zh or en)
            else:
                body = zh or en

            if not body:
                continue
            parts.append(f"{pos} {body}".strip())
        return "； ".join(parts)

    zh = trim_meaning(item.get("first_zh") or "")
    en = trim_meaning(item.get("first_en") or "")
    if mode == "en":
        return en
    if mode == "bilingual":
        return f"{zh} / {en}" if zh and en else (zh or en)
    if mode == "word":
        return ""
    return zh or en


class WeiciParser:
    title_re = re.compile(r"^#\s*\*\*\*\\#([^*]+)\*\*\*(?:[ \t]+(.+?))?[ \t]*$", re.M)
    uk_re = re.compile(r"^英音(?:[ \t]+([^\r\n<]+?))?[ \t]*$", re.M)
    us_re = re.compile(r"^美音(?:[ \t]+([^\r\n<]+?))?[ \t]*$", re.M)
    first_sense_re = re.compile(r"^###\s*1\.(.*)$")
    bold_text_re = re.compile(r"\*\*([^*]+)\*\*")
    sense_en_re = re.compile(r"^英译\s+(.+?)\s*$")
    audio_src_re = re.compile(r'<audio src="([^"]+)"')

    @classmethod
    def parse_markdown(cls, md_text: str, source_path: str) -> WordEntry | None:
        title_match = cls.title_re.search(md_text)
        if not title_match:
            return None

        word = title_match.group(1).strip()
        pos = (title_match.group(2) or "").strip()

        ipa_uk = cls.extract_ipa(md_text, cls.uk_re)
        ipa_us = cls.extract_ipa(md_text, cls.us_re)
        audio_uk = cls.extract_audio(md_text, "英音")
        audio_us = cls.extract_audio(md_text, "美音")
        first_zh, first_en = cls.extract_first_sense(md_text)

        if not first_zh:
            return None

        return WordEntry(
            word=word,
            pos=pos,
            category=build_category(word, source_path),
            ipa_uk=ipa_uk,
            ipa_us=ipa_us,
            audio_uk=audio_uk,
            audio_us=audio_us,
            first_zh=first_zh,
            first_en=first_en,
            source_path=source_path,
        )

    @staticmethod
    def extract_ipa(md_text: str, pattern: re.Pattern[str]) -> str:
        for match in pattern.finditer(md_text):
            value = (match.group(1) or "").strip()
            if value:
                return value
        return ""

    @classmethod
    def extract_audio(cls, md_text: str, label: str) -> str:
        lines = md_text.splitlines()
        for idx, raw_line in enumerate(lines):
            if raw_line.strip().startswith(label):
                for follow_line in lines[idx + 1 : idx + 6]:
                    match = cls.audio_src_re.search(follow_line)
                    if match:
                        return build_audio_url(match.group(1))
        return ""

    @classmethod
    def extract_first_sense(cls, md_text: str) -> tuple[str, str]:
        first_zh = ""
        first_en = ""
        in_first_sense = False

        for raw_line in md_text.splitlines():
            line = raw_line.strip()

            if not in_first_sense:
                sense_match = cls.first_sense_re.match(line)
                if not sense_match:
                    continue

                in_first_sense = True
                bold_parts = cls.bold_text_re.findall(line)
                if bold_parts:
                    first_zh = bold_parts[-1].strip()
                else:
                    first_zh = sense_match.group(1).strip().strip("*").strip()
                continue

            if re.match(r"^###\s*\d+\.", line) or re.match(r"^##\s+", line) or re.match(r"^#\s+", line):
                break

            en_match = cls.sense_en_re.match(line)
            if en_match:
                first_en = en_match.group(1).strip()
                break

        return first_zh, first_en

    @classmethod
    def parse_file(cls, md_text: str, source_path: str) -> List[WordEntry]:
        starts = list(re.finditer(r"^#\s*\*\*\*\\#.+$", md_text, re.M))
        if not starts:
            return []

        entries: List[WordEntry] = []
        for idx, start in enumerate(starts):
            section_start = start.start()
            section_end = starts[idx + 1].start() if idx + 1 < len(starts) else len(md_text)
            section = md_text[section_start:section_end]
            item = cls.parse_markdown(section, source_path)
            if item:
                entries.append(item)
        return entries


def _download_one(file_name: str) -> tuple[str, str]:
    url = f"{RAW_BASE}/{file_name}"
    last_error = None
    for attempt in range(3):
        try:
            resp = requests.get(url, timeout=40)
            resp.raise_for_status()
            return file_name, resp.text.replace("\r\n", "\n").replace("\r", "\n")
        except requests.RequestException as exc:
            last_error = exc
            if attempt < 2:
                time.sleep(1.2 * (attempt + 1))
    raise last_error


def build_category(word: str, source_path: str) -> str:
    if source_path == "weici_phrase.md":
        return "phrase"

    first_char = word[:1].upper()
    if "A" <= first_char <= "Z":
        return first_char
    return "special"


def build_audio_url(src: str) -> str:
    cleaned = (src or "").strip()
    if not cleaned:
        return ""
    if cleaned.startswith("http://") or cleaned.startswith("https://"):
        parsed = urlparse(cleaned)
        if parsed.netloc == "raw.githubusercontent.com" and parsed.path.startswith("/1299172402/weici/master/docs/2/media/"):
            media_name = quote(unquote(parsed.path.rsplit("/", 1)[-1]), safe="")
            return f"{PAGES_MEDIA_BASE}/{media_name}"
        if parsed.netloc == "1299172402.github.io" and parsed.path.startswith("/weici/media/"):
            media_name = quote(unquote(parsed.path.rsplit("/", 1)[-1]), safe="")
            return f"{PAGES_MEDIA_BASE}/{media_name}"
        encoded_path = quote(unquote(parsed.path), safe="/")
        return parsed._replace(path=encoded_path).geturl()
    if cleaned.startswith("./"):
        cleaned = cleaned[2:]
    if cleaned.startswith("media/"):
        media_name = quote(unquote(cleaned.rsplit("/", 1)[-1]), safe="")
        return f"{PAGES_MEDIA_BASE}/{media_name}"
    encoded = quote(cleaned, safe="/")
    return f"{RAW_BASE}/{encoded}"


def guess_audio_mimetype(url: str) -> str:
    path = urlparse(url).path.lower()
    if path.endswith(".mp3"):
        return "audio/mpeg"
    if path.endswith(".m4a"):
        return "audio/mp4"
    if path.endswith(".aac"):
        return "audio/aac"
    if path.endswith(".wav"):
        return "audio/wav"
    if path.endswith(".ogg"):
        return "audio/ogg"
    return "application/octet-stream"


def sanitize_audio_source(url: str) -> str | None:
    if not url:
        return None
    parsed = urlparse(url.strip())
    if parsed.scheme not in {"http", "https"}:
        return None
    if parsed.netloc == "raw.githubusercontent.com" and parsed.path.startswith("/1299172402/weici/master/docs/2/media/"):
        media_name = quote(unquote(parsed.path.rsplit("/", 1)[-1]), safe="")
        return f"{PAGES_MEDIA_BASE}/{media_name}"
    if parsed.netloc == "1299172402.github.io" and parsed.path.startswith("/weici/media/"):
        media_name = quote(unquote(parsed.path.rsplit("/", 1)[-1]), safe="")
        return f"{PAGES_MEDIA_BASE}/{media_name}"
    if parsed.netloc != "raw.githubusercontent.com":
        return None
    if not parsed.path.startswith("/1299172402/weici/master/docs/2/"):
        return None
    encoded_path = quote(unquote(parsed.path), safe="/")
    return parsed._replace(path=encoded_path).geturl()


def _download_and_parse() -> List[WordEntry]:
    DATA_DIR.mkdir(parents=True, exist_ok=True)

    words: List[WordEntry] = []
    with ThreadPoolExecutor(max_workers=4) as pool:
        futures = [pool.submit(_download_one, name) for name in SOURCE_FILES]
        for fut in as_completed(futures):
            file_name, text = fut.result()
            words.extend(WeiciParser.parse_file(text, file_name))

    words.sort(key=lambda x: x.word.lower())
    return words


def _write_cache(words: List[WordEntry]) -> None:
    payload = {
        "updated_at": datetime.now().isoformat(timespec="seconds"),
        "count": len(words),
        "words": [asdict(x) for x in words],
    }
    CACHE_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _load_notebook() -> list[dict]:
    NOTEBOOK_DIR.mkdir(parents=True, exist_ok=True)
    if not NOTEBOOK_PATH.exists():
        NOTEBOOK_PATH.write_text(json.dumps({"entries": []}, ensure_ascii=False, indent=2), encoding="utf-8")
        return []

    data = json.loads(NOTEBOOK_PATH.read_text(encoding="utf-8"))
    return data.get("entries", [])


def _write_notebook(entries: list[dict]) -> None:
    NOTEBOOK_DIR.mkdir(parents=True, exist_ok=True)
    payload = {
        "updated_at": datetime.now().isoformat(timespec="seconds"),
        "count": len(entries),
        "entries": entries,
    }
    NOTEBOOK_PATH.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _load_cache() -> List[WordEntry]:
    if not CACHE_PATH.exists():
        return []
    data = json.loads(CACHE_PATH.read_text(encoding="utf-8"))
    items = []
    for x in data.get("words", []):
        if "category" not in x:
            x["category"] = build_category(x.get("word", ""), x.get("source_path", ""))
        x.setdefault("audio_uk", "")
        x.setdefault("audio_us", "")
        items.append(WordEntry(**x))
    return items


def get_words(force_refresh: bool = False) -> List[WordEntry]:
    if force_refresh or not CACHE_PATH.exists():
        words = _download_and_parse()
        _write_cache(words)
        return words
    return _load_cache()


@app.get("/")
def home():
    return render_template("index.html")


@app.get("/api/words")
def api_words():
    refresh = request.args.get("refresh", "0") == "1"
    words = get_words(force_refresh=refresh)
    aggregated = aggregate_words(words)
    return jsonify(
        {
            "count": len(aggregated),
            "raw_count": len(words),
            "updated_at": datetime.fromtimestamp(CACHE_PATH.stat().st_mtime).isoformat(timespec="seconds")
            if CACHE_PATH.exists()
            else "",
            "words": aggregated,
        }
    )


@app.get("/api/notebook")
def api_notebook():
    entries = _load_notebook()
    return jsonify({"count": len(entries), "entries": entries})


@app.post("/api/notebook")
def api_notebook_save():
    payload = request.get_json(force=True, silent=True) or {}
    entries = payload.get("entries", [])
    cleaned = [item for item in entries if isinstance(item, dict) and item.get("word")]
    _write_notebook(cleaned)
    return jsonify({"count": len(cleaned), "entries": cleaned})


@app.post("/api/export")
def api_export():
    payload = request.get_json(force=True, silent=True) or {}
    entries = payload.get("entries", [])
    mode = payload.get("mode", "zh")  # zh | bilingual | en
    file_type = payload.get("file_type", "docx")  # docx | txt

    if not entries:
        return jsonify({"error": "生词表为空，请先添加单词。"}), 400

    rows = []
    for i, item in enumerate(entries, start=1):
        word = (item.get("word") or "").strip()
        ipa = (item.get("ipa") or item.get("ipa_us") or item.get("ipa_uk") or "").strip()
        pos = (item.get("pos") or "").strip()

        if not word:
            continue

        meaning = build_export_meaning(item, mode)

        line = f"{i}. {word}"
        if pos:
            line += f" ({pos})"
        if ipa:
            line += f" [{ipa}]"
        if meaning:
            line += f" - {meaning}"
        rows.append(line)

    if file_type == "txt":
        content = "\n".join(rows)
        bytes_io = io.BytesIO(content.encode("utf-8"))
        filename = f"weici_vocab_{mode}.txt"
        mime = "text/plain"
    else:
        if Document is None:
            return jsonify({"error": "未安装 python-docx，无法导出 docx。请先 pip install -r requirements.txt"}), 500

        doc = Document()
        doc.add_heading("维词生词本", level=1)
        doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        for row in rows:
            doc.add_paragraph(row)

        bytes_io = io.BytesIO()
        doc.save(bytes_io)
        bytes_io.seek(0)
        filename = f"weici_vocab_{mode}.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return send_file(
        bytes_io,
        as_attachment=True,
        download_name=filename,
        mimetype=mime,
    )


@app.get("/api/audio")
def api_audio():
    src = request.args.get("src", "")
    url = sanitize_audio_source(src)
    if not url:
        return jsonify({"error": "无效的音频地址"}), 400

    try:
        upstream = requests.get(url, timeout=30)
        upstream.raise_for_status()
    except requests.RequestException:
        return jsonify({"error": "音频获取失败"}), 502

    mimetype = upstream.headers.get("Content-Type") or guess_audio_mimetype(url)
    return Response(
        upstream.content,
        mimetype=mimetype,
        headers={
            "Cache-Control": "public, max-age=86400",
            "Accept-Ranges": "bytes",
        },
    )


if __name__ == "__main__":
    debug = os.environ.get("WEICI_DEBUG", "1") == "1"
    host = os.environ.get("WEICI_HOST", "127.0.0.1")
    port = int(os.environ.get("PORT", "5050"))
    app.run(host=host, port=port, debug=debug, use_reloader=debug)
